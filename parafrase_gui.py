
import streamlit as st
from docx import Document
import fitz  # PyMuPDF
from transformers import pipeline
from fuzzywuzzy import fuzz
import tempfile
import os

@st.cache_resource
def load_paraphraser():
    return pipeline("text2text-generation", model="Vamsi/T5_Paraphrase_Paws")

paraphraser = load_paraphraser()

def load_docx_text(docx_file):
    doc = Document(docx_file)
    paragraphs = [p.text for p in doc.paragraphs]
    return paragraphs, doc

def load_pdf_text(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def detect_similar_paragraphs(paragraphs, turnitin_text, threshold=70):
    flagged = []
    for i, para in enumerate(paragraphs):
        score = fuzz.partial_ratio(para, turnitin_text)
        if score > threshold and len(para.strip()) > 30:
            flagged.append((i, para, score))
    return flagged

def academic_paraphrase(text):
    result = paraphraser(f"paraphrase: {text} </s>", max_length=256, do_sample=True,
                         top_k=120, top_p=0.95, num_return_sequences=1)
    return result[0]['generated_text']

def replace_paragraphs(doc, flagged_indices, new_paragraphs):
    for (i, _, _), new_text in zip(flagged_indices, new_paragraphs):
        doc.paragraphs[i].text = new_text
    return doc

def main():
    st.set_page_config(page_title="Parafrase Skripsi Akademik", layout="centered")
    st.title("📘 Mesin Parafrase Akademik")
    st.markdown("Upload skripsi (.docx) dan hasil Turnitin (.pdf) untuk diproses.")

    docx_file = st.file_uploader("📄 Upload File Skripsi (.docx)", type=["docx"])
    pdf_file = st.file_uploader("📑 Upload Hasil Turnitin (.pdf)", type=["pdf"])

    threshold = st.slider("🎯 Sensitivitas Deteksi Plagiarisme (%)", 50, 100, 70)

    if st.button("🚀 Mulai Parafrase") and docx_file and pdf_file:
        with st.spinner("Memproses dokumen..."):
            paras, doc = load_docx_text(docx_file)
            turnitin_text = load_pdf_text(pdf_file)
            flagged = detect_similar_paragraphs(paras, turnitin_text, threshold=threshold)

            if not flagged:
                st.success("✅ Tidak ditemukan paragraf yang terdeteksi plagiasi/AI.")
                return

            st.info(f"🧠 Ditemukan {len(flagged)} paragraf yang perlu diparafrase.")

            new_paras = []
            for i, (idx, para, score) in enumerate(flagged):
                st.write(f"---\n🔍 Paragraf {idx+1} (Kemiripan: {score}%)")
                st.text_area("Teks Asli", para, height=100, key=f"ori_{i}")
                paraphrased = academic_paraphrase(para)
                st.text_area("Hasil Parafrase", paraphrased, height=100, key=f"parafrase_{i}")
                new_paras.append(paraphrased)

            doc = replace_paragraphs(doc, flagged, new_paras)

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                temp_output_path = tmp.name
                doc.save(temp_output_path)

            with open(temp_output_path, "rb") as f:
                st.download_button(
                    label="💾 Download Hasil Parafrase",
                    data=f,
                    file_name="skripsi_parafrase.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            os.remove(temp_output_path)

if __name__ == "__main__":
    main()
